from __future__ import annotations

from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / 'report' / 'Vinci_Technical_Report_No_2_unstyled.docx'
OUT = ROOT / 'report' / 'Vinci_Technical_Report_No_2.docx'

TEAL = '08756F'
DARK = '1C2628'
GRAY = '6D7375'
PALE = 'E6F1EF'
LIGHT_GRAY = 'F4F5F5'
WHITE = 'FFFFFF'
ORANGE = 'C96A12'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in edges:
            tag = f'w:{edge}'
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edges[edge].items():
                element.set(qn(f'w:{key}'), str(value))


def set_paragraph_bottom_border(paragraph, color: str = TEAL, size: str = '10') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        p_bdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        p_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_run_font(run, name: str, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_all_runs(paragraph, name: str, size: float, color: str = DARK) -> None:
    for run in paragraph.runs:
        set_run_font(run, name, size, color)


def replace_long_hex_with_breaks(paragraph) -> None:
    # Word does not reliably line-break 64-character hashes. Inserting zero-width
    # break opportunities keeps the visible identifier unchanged while avoiding overflow.
    for run in paragraph.runs:
        if not re.search(r'[0-9a-f]{48,}', run.text):
            continue
        run.text = re.sub(
            r'([0-9a-f]{48,})',
            lambda m: '\u200b'.join(m.group(1)[i:i+8] for i in range(0, len(m.group(1)), 8)),
            run.text,
        )


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    for style_name in ('Normal', 'Body Text', 'First Paragraph', 'Compact', 'Block Text'):
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = 'Liberation Serif'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Liberation Serif')
        style.font.size = Pt(10.2)
        style.font.color.rgb = RGBColor.from_string(DARK)
        pf = style.paragraph_format
        pf.space_after = Pt(3.0)
        pf.line_spacing = 1.04
        pf.widow_control = True

    if 'Body Text' in styles:
        styles['Body Text'].paragraph_format.first_line_indent = Inches(0.18)
        styles['Body Text'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if 'First Paragraph' in styles:
        styles['First Paragraph'].paragraph_format.first_line_indent = Inches(0)
        styles['First Paragraph'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if 'Compact' in styles:
        styles['Compact'].paragraph_format.left_indent = Inches(0.22)
        styles['Compact'].paragraph_format.first_line_indent = Inches(-0.12)
        styles['Compact'].paragraph_format.space_after = Pt(1)
    if 'Block Text' in styles:
        styles['Block Text'].paragraph_format.left_indent = Inches(0.35)
        styles['Block Text'].paragraph_format.right_indent = Inches(0.35)
        styles['Block Text'].paragraph_format.space_before = Pt(5)
        styles['Block Text'].paragraph_format.space_after = Pt(5)
        styles['Block Text'].font.italic = True

    # Pandoc maps Markdown h1/h2/h3 to Word Heading 1/2/3.
    h1 = next(st for st in styles if st.style_id == 'Heading1')
    h1.font.name = 'Liberation Sans'
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Liberation Sans')
    h1.font.size = Pt(25)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(DARK)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = next(st for st in styles if st.style_id == 'Heading2')
    h2.font.name = 'Liberation Sans'
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Liberation Sans')
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK)
    h2.paragraph_format.space_before = Pt(13)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = next(st for st in styles if st.style_id == 'Heading3')
    h3.font.name = 'Liberation Sans'
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Liberation Sans')
    h3.font.size = Pt(11.3)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    for style_name in ('Image Caption', 'Caption'):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = 'Liberation Serif'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Liberation Serif')
            style.font.size = Pt(8.5)
            style.font.color.rgb = RGBColor.from_string(GRAY)
            style.paragraph_format.space_before = Pt(2)
            style.paragraph_format.space_after = Pt(7)
            style.paragraph_format.keep_with_next = False


def configure_section(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.62)
        section.bottom_margin = Inches(0.62)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.header_distance = Inches(0.25)
        section.footer_distance = Inches(0.28)

        header = section.header
        header.is_linked_to_previous = False
        table = header.add_table(rows=1, cols=2, width=Inches(7.0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(3.5)
        table.columns[1].width = Inches(3.5)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        left.text = 'VINCI RESEARCH  •  TECHNICAL REPORT NO. 2'
        right.text = 'CHARACTER TRANSFER ACROSS MODEL FAMILIES'
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, 'Liberation Sans', 7.2, DARK)
            set_cell_border(cell, bottom={'val': 'single', 'sz': '8', 'color': TEAL})
        # Remove the empty default paragraph above the table from visual flow.
        if header.paragraphs:
            header.paragraphs[0].text = ''
            header.paragraphs[0].paragraph_format.space_after = Pt(0)
            header.paragraphs[0].paragraph_format.line_spacing = Pt(1)

        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.add_run('VINCI TECHNICAL REPORT NO. 2   •   ')
        add_page_field(p)
        for r in p.runs:
            set_run_font(r, 'Liberation Sans', 7.2, GRAY)


def style_title_block(doc: Document) -> None:
    p0 = doc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(12)
    p0.paragraph_format.space_after = Pt(2)
    style_all_runs(p0, 'Liberation Sans', 25, DARK)
    for r in p0.runs:
        r.bold = True
    set_paragraph_bottom_border(p0, TEAL, '12')

    p1 = doc.paragraphs[1]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(10)
    style_all_runs(p1, 'Liberation Sans', 13.2, GRAY)
    for r in p1.runs:
        r.bold = False

    p2 = doc.paragraphs[2]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(5)
    p2.paragraph_format.keep_with_next = True
    style_all_runs(p2, 'Liberation Sans', 10.2, DARK)
    for r in p2.runs:
        r.bold = True

    p3 = doc.paragraphs[3]
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.left_indent = Inches(0.12)
    p3.paragraph_format.right_indent = Inches(0.12)
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(10)
    p3.paragraph_format.keep_with_next = True
    style_all_runs(p3, 'Liberation Sans', 8.8, DARK)
    for r in p3.runs:
        r.bold = True
    set_paragraph_shading(p3, PALE)

    # Abstract heading: visually distinct without starting a new page.
    p4 = doc.paragraphs[4]
    p4.paragraph_format.space_before = Pt(3)
    p4.paragraph_format.space_after = Pt(3)
    style_all_runs(p4, 'Liberation Sans', 13, TEAL)
    for r in p4.runs:
        r.bold = True

    # Keywords line.
    p9 = doc.paragraphs[9]
    p9.paragraph_format.space_before = Pt(4)
    p9.paragraph_format.space_after = Pt(5)
    style_all_runs(p9, 'Liberation Serif', 9.2, DARK)


def style_figures(doc: Document) -> None:
    max_width = Inches(6.85)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Captioned Figure':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
        elif p.style.name == 'Image Caption':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = False
            for r in p.runs:
                set_run_font(r, 'Liberation Serif', 8.5, GRAY)
        replace_long_hex_with_breaks(p)


def style_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.style = 'Table Grid' if 'Table Grid' in [s.name for s in doc.styles] else table.style
        for ri, row in enumerate(table.rows):
            keep_row_together(row)
            if ri == 0:
                set_repeat_table_header(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.margin_top = Inches(0.03)
                cell.margin_bottom = Inches(0.03)
                cell.margin_left = Inches(0.04)
                cell.margin_right = Inches(0.04)
                if ri == 0:
                    set_cell_shading(cell, PALE)
                    set_cell_border(
                        cell,
                        top={'val': 'single', 'sz': '8', 'color': TEAL},
                        bottom={'val': 'single', 'sz': '8', 'color': TEAL},
                        left={'val': 'nil'}, right={'val': 'nil'}
                    )
                else:
                    set_cell_border(
                        cell,
                        top={'val': 'nil'},
                        bottom={'val': 'single', 'sz': '3', 'color': 'D5D9DA'},
                        left={'val': 'nil'}, right={'val': 'nil'}
                    )
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.keep_together = True
                    replace_long_hex_with_breaks(p)
                    for r in p.runs:
                        set_run_font(r, 'Liberation Sans' if ri == 0 else 'Liberation Serif', 8.0 if ri == 0 else 8.3, DARK, bold=(True if ri == 0 else None))


def style_body(doc: Document) -> None:
    in_references = False
    in_appendix = False
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if idx <= 9:
            continue
        if text == 'References':
            in_references = True
        elif text.startswith('Appendix A.'):
            in_references = False
            in_appendix = True
            p.paragraph_format.page_break_before = True
        elif text.startswith('Appendix ') and in_appendix:
            p.paragraph_format.page_break_before = False

        replace_long_hex_with_breaks(p)

        if p.style.name in ('Body Text', 'First Paragraph'):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if in_references and p.style.name not in ('Heading 2', 'Heading 3'):
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.22)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                set_run_font(r, 'Liberation Serif', 8.8, DARK)
        if in_appendix and p.style.name in ('Heading 2', 'Heading 3'):
            for r in p.runs:
                set_run_font(r, 'Liberation Sans', 13 if p.style.name == 'Heading 2' else 10.8, DARK, bold=True)
        if in_appendix and p.style.name == 'Compact':
            p.paragraph_format.space_after = Pt(0.5)
            p.paragraph_format.line_spacing = 0.95
            for r in p.runs:
                set_run_font(r, 'Liberation Serif', 9.1, DARK)

    # The experimental-design figure naturally rolls to page 2. Let the introduction
    # follow it instead of forcing an otherwise mostly blank front-matter page.


def main() -> None:
    doc = Document(SRC)
    configure_styles(doc)
    configure_section(doc)
    style_title_block(doc)
    style_figures(doc)
    style_tables(doc)
    style_body(doc)

    cp = doc.core_properties
    cp.title = 'Character Transfer Across Three Model Families'
    cp.subject = 'Vinci Technical Report No. 2 - development-tier character-transfer study'
    cp.author = 'George Pu; Ayush Naik'
    cp.keywords = 'character post-training; DPO; unsupported assertions; answer preservation; evaluation reliability'
    cp.category = 'Vinci Research Technical Report'
    cp.comments = 'Version 0.9 publication draft. Internal review only.'

    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    main()


# --- Strip the template's stale application statistics -----------------------
# python-docx inherits docProps/app.xml from its starter template, which reports
# the template's page/word counts and "Microsoft Word 12.0.0" rather than this
# document's. Word and LibreOffice recompute on open, but the distributed file
# should not carry values that describe a different document.
def _strip_app_properties(docx_path):
    import shutil, zipfile
    from pathlib import Path
    src = Path(docx_path)
    tmp = src.with_suffix('.stripped.docx')
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'docProps/app.xml':
                data = (b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                        b'<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/'
                        b'2006/extended-properties"><Application>Vinci Research publication '
                        b'build</Application></Properties>')
            zout.writestr(item, data)
    shutil.move(str(tmp), str(src))


_strip_app_properties(OUT)
